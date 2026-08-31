import json
from pathlib import Path

import pymupdf
import pytest
from docx import Document as DocxDocument

from licita_core.r2_skeleton import esboco_processo, main
from licita_core.schema import DocumentType, ProcurementProcess


def _pdf(path: Path, linhas: list[str]) -> None:
    document = pymupdf.open()
    page = document.new_page(width=480, height=300)
    y = 42
    for linha in linhas:
        page.insert_text((40, y), linha, fontsize=11)
        y += 18
    document.save(path)
    document.close()


def _docx(path: Path) -> None:
    document = DocxDocument()
    document.add_paragraph("Estudo técnico preliminar sintético.")
    document.add_paragraph("Quantidade estimada: 12 unidades.")
    document.save(path)


def test_esboco_preserva_blocos_e_valida_schema(tmp_path: Path) -> None:
    etp = tmp_path / "etp.docx"
    tr = tmp_path / "tr.pdf"
    _docx(etp)
    _pdf(tr, ["TERMO DE REFERÊNCIA", "Item 1: caneta azul", "Quantidade: 12 un"])

    processo = esboco_processo(
        "proc-r2-esboco",
        [(etp, DocumentType.ETP), (tr, DocumentType.TR)],
    )
    roundtrip = ProcurementProcess.model_validate(processo.model_dump(mode="json"))

    assert roundtrip.id == "proc-r2-esboco"
    assert [doc.type for doc in roundtrip.documents] == [
        DocumentType.ETP,
        DocumentType.TR,
    ]
    assert roundtrip.documents[0].items == []
    assert roundtrip.documents[1].field_values == []
    tr_text = " ".join(
        block.text
        for section in roundtrip.documents[1].sections
        for block in section.blocks
    )
    assert "caneta azul" in tr_text
    for document in roundtrip.documents:
        for section in document.sections:
            ids = {block.id for block in section.blocks}
            assert section.evidence.block_id in ids
            bloco = next(
                block for block in section.blocks if block.id == section.evidence.block_id
            )
            assert section.evidence.quote in bloco.text


def test_cli_escreve_json(tmp_path: Path) -> None:
    etp = tmp_path / "etp.docx"
    tr = tmp_path / "tr.pdf"
    saida = tmp_path / "proc.json"
    _docx(etp)
    _pdf(tr, ["TR sintético com texto suficiente."])

    assert main(["--id", "p1", "--etp", str(etp), "--tr", str(tr), "--saida", str(saida)]) == 0
    payload = json.loads(saida.read_text(encoding="utf-8"))
    assert payload["id"] == "p1"
    assert len(payload["documents"]) == 2


def test_cli_exige_documento(tmp_path: Path) -> None:
    with pytest.raises(SystemExit):
        main(["--id", "p1", "--saida", str(tmp_path / "x.json")])
