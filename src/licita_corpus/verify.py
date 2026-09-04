"""Verificação local de documentos e fallback opcional de OCR.

A verificação normal continua sendo somente leitura: ``verificar(path)`` abre o
arquivo, extrai o texto e nunca procura nem executa o Tesseract. O OCR é um
fallback explícito para páginas de PDF que não têm texto suficiente.

O fallback é deliberadamente limitado. Cada página é rasterizada em um
arquivo temporário, o Tesseract produz TSV em disco e o resultado só é lido
após um limite de bytes. O PDF original nunca é salvo novamente ou alterado.
"""

from __future__ import annotations

import csv
import hashlib
import io
import math
import os
import shutil
import stat
import subprocess
import tempfile
import time
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Mapping

#: Critérios mínimos para considerar uma página utilizável.
MIN_CARACTERES_POR_PAGINA = 40
MIN_PALAVRAS_POR_PAGINA = 8
MIN_CONFIANCA_OCR = 40.0
MAX_FRACAO_CARACTERES_CONTROLE = 0.05
MIN_CARACTERES_CONTROLE_CORROMPIDOS = 8

#: Limites seguros padrão do fallback OCR. São argumentos configuráveis de
#: :func:`verificar` e também ficam expostos para consumidores do módulo.
MAX_PAGINAS_OCR = 100
MAX_OCR_PAGES = MAX_PAGINAS_OCR
MAX_TEMPO_OCR_DOCUMENTO = 300.0
MAX_OCR_TIMEOUT_TOTAL = MAX_TEMPO_OCR_DOCUMENTO
MAX_TEMPO_OCR_PAGINA = 30.0
MAX_OCR_TIMEOUT_PAGINA = MAX_TEMPO_OCR_PAGINA
DPI_OCR = 200
MAX_MEGAPIXELS_OCR = 20.0
MAX_OCR_MEGAPIXELS = MAX_MEGAPIXELS_OCR
MAX_BYTES_SAIDA_OCR = 2 * 1024 * 1024
MAX_OCR_OUTPUT_BYTES = MAX_BYTES_SAIDA_OCR

# Aliases em inglês/forma alternativa, para evitar que a configuração dependa
# do idioma do consumidor da biblioteca.
OCR_MAX_PAGINAS = MAX_PAGINAS_OCR
OCR_MAX_PAGES = MAX_PAGINAS_OCR
OCR_TIMEOUT_TOTAL = MAX_TEMPO_OCR_DOCUMENTO
OCR_TOTAL_TIMEOUT = MAX_TEMPO_OCR_DOCUMENTO
OCR_TOTAL_TIMEOUT_SECONDS = MAX_TEMPO_OCR_DOCUMENTO
OCR_TIMEOUT_PAGINA = MAX_TEMPO_OCR_PAGINA
OCR_PAGE_TIMEOUT = MAX_TEMPO_OCR_PAGINA
OCR_PAGE_TIMEOUT_SECONDS = MAX_TEMPO_OCR_PAGINA
OCR_DPI = DPI_OCR
OCR_MAX_MEGAPIXELS = MAX_MEGAPIXELS_OCR
OCR_MAX_PIXELS = int(MAX_MEGAPIXELS_OCR * 1_000_000)
OCR_MAX_RASTER_PIXELS = OCR_MAX_PIXELS
OCR_MAX_OUTPUT_BYTES = MAX_BYTES_SAIDA_OCR
OCR_OUTPUT_MAX_BYTES = MAX_BYTES_SAIDA_OCR


def monotonic() -> float:
    """Relógio monotônico patchável nos testes."""
    return time.monotonic()


@dataclass(frozen=True, slots=True)
class AvaliacaoPagina:
    """Texto efetivo e decisão de OCR para uma página física de PDF."""

    pagina: int
    caracteres: int
    precisa_ocr: bool
    ocr_usado: bool = False
    caracteres_originais: int = 0
    texto: str = field(default="", repr=False)
    texto_original: str = field(default="", repr=False)
    palavras: int = 0
    palavras_originais: int = 0
    # Quando houve uma tentativa que produziu TSV, a confiança fica registrada
    # mesmo que o texto tenha sido rejeitado e o original tenha sido mantido.
    confianca_media: float | None = None
    erro_ocr: str | None = None

    @property
    def confianca(self) -> float | None:
        """Alias curto para consumidores que não precisam do nome completo."""
        return self.confianca_media


@dataclass(slots=True)
class Verificacao:
    caminho: Path
    abriu: bool
    paginas: int | None = None
    caracteres: int = 0
    precisa_ocr: bool = False
    erro: str | None = None
    texto: str = field(default="", repr=False)
    # Os campos abaixo são aditivos: os sete campos anteriores mantêm a forma
    # usada pelo coletor e por consumidores que constroem Verificacao.
    ocr_solicitado: bool = False
    ocr_usado: bool = False
    ocr_motor: str | None = None
    ocr_idioma: str | None = None
    paginas_ocr: tuple[int, ...] = ()
    paginas_avaliadas: tuple[AvaliacaoPagina, ...] = ()
    texto_original: str = field(default="", repr=False)
    # A confiança é agregada de todos os TSVs lidos, inclusive os que não
    # foram aceitos. Assim uma rejeição também é auditável.
    ocr_confianca_media: float | None = None
    ocr_erros: tuple[str, ...] = ()
    paginas_ocr_tentadas: tuple[int, ...] = ()

    @property
    def utilizavel(self) -> bool:
        """Abre e entrega texto suficiente, por extração direta ou por OCR."""
        return self.abriu and not self.precisa_ocr and self.caracteres > 0

    @property
    def texto_por_pagina(self) -> tuple[str, ...]:
        """Texto efetivo de cada página, na ordem física do PDF."""
        return tuple(pagina.texto for pagina in self.paginas_avaliadas)

    @property
    def texto_original_por_pagina(self) -> tuple[str, ...]:
        """Texto original do PyMuPDF, sem substituição pelo OCR."""
        return tuple(pagina.texto_original for pagina in self.paginas_avaliadas)

    @property
    def caracteres_por_pagina(self) -> tuple[int, ...]:
        """Quantidade de caracteres efetivos em cada página."""
        return tuple(pagina.caracteres for pagina in self.paginas_avaliadas)

    @property
    def confianca_media_ocr(self) -> float | None:
        """Alias legível para a confiança agregada do OCR."""
        return self.ocr_confianca_media

    @property
    def ocr_confianca(self) -> float | None:
        """Alias compatível para a confiança agregada do OCR."""
        return self.ocr_confianca_media

    @property
    def confianca_ocr(self) -> float | None:
        """Alias compatível usado por alguns catálogos legados."""
        return self.ocr_confianca_media

    @property
    def ocr(self) -> dict[str, Any]:
        """Metadados do OCR em formato pronto para catálogo/JSON."""
        return {
            "solicitado": self.ocr_solicitado,
            "usado": self.ocr_usado,
            "motor": self.ocr_motor,
            "idioma": self.ocr_idioma,
            "paginas": list(self.paginas_ocr),
            "paginas_tentadas": list(self.paginas_ocr_tentadas),
            "confianca_media": self.ocr_confianca_media,
            "erros": list(self.ocr_erros),
        }


class _OcrErro(RuntimeError):
    """Falha operacional ou rejeição controlada do fallback OCR."""


@dataclass(frozen=True, slots=True)
class _QualidadeTexto:
    caracteres: int
    palavras: int
    confianca_media: float | None
    passou: bool
    motivos: tuple[str, ...]


@dataclass(frozen=True, slots=True)
class _ResultadoOcrPagina:
    texto: str
    palavras: int
    confianca_media: float | None


@dataclass(frozen=True, slots=True)
class _ResultadoOcrDocumento:
    paginas_aplicadas: frozenset[int]
    paginas_tentadas: frozenset[int]
    confiancas: Mapping[int, float | None]
    erros: tuple[str, ...]
    palavras: Mapping[int, int] = field(default_factory=dict)
    erros_por_pagina: Mapping[int, str] = field(default_factory=dict)


def _idioma_ocr(idioma: object) -> tuple[str, tuple[str, ...]]:
    if not isinstance(idioma, str) or not idioma.strip():
        raise _OcrErro("idioma OCR não pode ser vazio")
    partes = tuple(parte.strip() for parte in idioma.strip().split("+"))
    if not partes or any(not parte for parte in partes):
        raise _OcrErro(
            f"idioma OCR inválido: {idioma.strip()!r}; use um código como 'por' ou 'por+eng'"
        )
    return "+".join(partes), partes


def _validar_tesseract(
    executavel: str,
    idioma: str,
    componentes: tuple[str, ...],
    *,
    timeout: float = 15.0,
    diretorio: Path | None = None,
    limite_saida: int = MAX_BYTES_SAIDA_OCR,
    prazo: float | None = None,
) -> None:
    """Confere o executável e os idiomas sem capturar saída ilimitada."""
    try:
        encontrado = shutil.which(executavel)
    except (OSError, TypeError) as exc:
        raise _OcrErro(
            f"não foi possível localizar o executável do Tesseract {executavel!r}: {exc}"
        ) from exc
    if encontrado is None:
        raise _OcrErro(
            f"executável do Tesseract não encontrado: {executavel!r}; "
            "instale o Tesseract ou desative o OCR"
        )

    if diretorio is None:
        with tempfile.TemporaryDirectory(prefix="licita-corpus-ocr-") as temporario:
            _validar_tesseract(
                executavel,
                idioma,
                componentes,
                timeout=timeout,
                diretorio=Path(temporario),
                limite_saida=limite_saida,
                prazo=prazo,
            )
        return

    stdout_path = diretorio / "idiomas.stdout"
    stderr_path = diretorio / "idiomas.stderr"
    prazo_validacao = prazo
    if prazo_validacao is None:
        prazo_validacao = monotonic() + timeout
    restante_validacao = prazo_validacao - monotonic()
    if restante_validacao <= 0:
        raise subprocess.TimeoutExpired([executavel, "--list-langs"], timeout)
    try:
        resultado = _rodar_subprocesso(
            [executavel, "--list-langs"],
            timeout=restante_validacao,
            stdout_path=stdout_path,
            stderr_path=stderr_path,
            limite_saida=limite_saida,
            prazo=prazo_validacao,
        )
    except subprocess.TimeoutExpired as exc:
        raise _OcrErro(
            f"TimeoutExpired: tempo total/consulta de idiomas do Tesseract esgotado "
            f"para {executavel!r} (idioma {idioma!r})"
        ) from exc
    except FileNotFoundError as exc:
        raise _OcrErro(
            f"executável do Tesseract não encontrado ao consultar idiomas: {executavel!r}"
        ) from exc
    except OSError as exc:
        raise _OcrErro(
            f"não foi possível executar o Tesseract {executavel!r}: {exc}"
        ) from exc
    except _OcrErro:
        raise
    except Exception as exc:
        raise _OcrErro(
            f"erro ao consultar os idiomas do Tesseract {executavel!r}: "
            f"{type(exc).__name__}: {exc}"
        ) from exc

    if getattr(resultado, "returncode", 1) != 0:
        detalhe = _detalhe_subprocesso(
            resultado,
            stderr_path,
            limite_saida,
            prazo=prazo_validacao,
        )
        sufixo = f": {detalhe}" if detalhe else ""
        raise _OcrErro(
            f"Tesseract retornou código não-zero ao consultar idiomas "
            f"(idioma {idioma!r}){sufixo}"
        )

    try:
        conteudo = _conteudo_limitado_com_prazo(
            stdout_path,
            getattr(resultado, "stdout", None),
            limite_saida,
            "a listagem de idiomas do Tesseract",
            prazo_validacao,
        )
    except _OcrErro:
        raise

    disponiveis = {
        linha.strip()
        for linha in conteudo.splitlines()
        if linha.strip()
        and not linha.lower().startswith("list of available languages")
    }
    faltantes = tuple(componente for componente in componentes if componente not in disponiveis)
    if faltantes:
        lista = ", ".join(sorted(disponiveis)) or "nenhum"
        raise _OcrErro(
            f"idioma(s) OCR não instalado(s) no Tesseract: {', '.join(faltantes)} "
            f"(solicitado: {idioma}); idiomas disponíveis: {lista}"
        )


_TAMANHO_CHUNK_OCR = 64 * 1024
_INTERVALO_MONITORAMENTO_OCR = 0.01


def _status_arquivo_regular(
    caminho: Path,
    contexto: str,
    *,
    ausente_ok: bool = True,
) -> os.stat_result | None:
    """Obtém o ``lstat`` e rejeita links e arquivos especiais."""
    try:
        status = os.lstat(caminho)
    except FileNotFoundError:
        if ausente_ok:
            return None
        raise _OcrErro(f"{contexto} não existe")
    except OSError as exc:
        raise _OcrErro(f"não foi possível verificar {contexto}: {exc}") from exc

    if not stat.S_ISREG(status.st_mode):
        if stat.S_ISLNK(status.st_mode):
            tipo = "symlink"
        else:
            tipo = "arquivo especial"
        raise _OcrErro(f"{contexto} deve ser arquivo regular; encontrado {tipo}")
    return status


def _opener_seguro(caminho: str | bytes, flags: int) -> int:
    # O_NONBLOCK é inócuo para arquivo regular e impede que uma troca por FIFO
    # bloqueie o processo pai antes de o fstat rejeitá-la.
    flags |= getattr(os, "O_NONBLOCK", 0)
    flags |= getattr(os, "O_NOFOLLOW", 0)
    flags |= getattr(os, "O_CLOEXEC", 0)
    return os.open(caminho, flags, 0o600)


def _abrir_saida_regular(caminho: Path, contexto: str) -> Any:
    """Abre uma saída sem seguir symlink nem bloquear em FIFO."""
    _status_arquivo_regular(caminho, contexto)
    try:
        # Usar open(opener=...) conserva ``.name`` para doubles de Popen e
        # ainda aplica as flags de segurança no descritor real.
        arquivo = open(
            caminho,
            "wb",
            buffering=0,
            opener=_opener_seguro,
        )
    except OSError as exc:
        raise _OcrErro(f"não foi possível criar {contexto}: {exc}") from exc

    try:
        status = os.fstat(arquivo.fileno())
        if not stat.S_ISREG(status.st_mode):
            raise _OcrErro(f"{contexto} deve ser arquivo regular")
        # Sem buffering, mocks e o processo filho tornam o crescimento físico
        # observável pelo monitor imediatamente.
        return arquivo
    except Exception:
        arquivo.close()
        raise


def _criar_saida_regular(caminho: Path, contexto: str) -> None:
    arquivo = _abrir_saida_regular(caminho, contexto)
    arquivo.close()


def _caminhos_unicos(caminhos: tuple[Path, ...] | list[Path]) -> tuple[Path, ...]:
    resultado: list[Path] = []
    vistos: set[Path] = set()
    for caminho in caminhos:
        caminho = Path(caminho)
        if caminho not in vistos:
            vistos.add(caminho)
            resultado.append(caminho)
    return tuple(resultado)


def _tamanho_saida_fisico(
    caminhos: tuple[Path, ...] | list[Path],
    *,
    arquivos_abertos: tuple[tuple[Path, Any], ...] = (),
) -> int:
    """Soma o tamanho físico das saídas, sem seguir links ou especiais."""
    total = 0
    caminhos_com_descritor = {Path(caminho) for caminho, _ in arquivos_abertos}
    for caminho in _caminhos_unicos(caminhos):
        caminho = Path(caminho)
        status = _status_arquivo_regular(caminho, f"a saída OCR {caminho}")
        if status is not None and caminho not in caminhos_com_descritor:
            total += max(0, int(status.st_size))

    # Se o nome for removido enquanto o filho ainda escreve, o fstat do
    # descritor continua medindo o arquivo físico que está recebendo a saída.
    for caminho, arquivo in arquivos_abertos:
        try:
            status = os.fstat(arquivo.fileno())
        except OSError as exc:
            raise _OcrErro(
                f"não foi possível verificar a saída OCR {caminho}: {exc}"
            ) from exc
        if not stat.S_ISREG(status.st_mode):
            raise _OcrErro(f"a saída OCR {caminho} deve ser arquivo regular")
        total += max(0, int(status.st_size))
    return total


def _matar_e_aguardar(processo: Any, *, forcar: bool = False) -> None:
    """Mata um processo ainda vivo e sempre chama ``wait``."""
    try:
        try:
            vivo = forcar or processo.poll() is None
        except Exception:
            vivo = True
        if vivo:
            try:
                processo.kill()
            except (OSError, ProcessLookupError):
                pass
    finally:
        # Mesmo se kill falhar, não deixamos um filho/zumbi pendurado.
        processo.wait()


def _rodar_subprocesso(
    argumentos: list[str],
    *,
    timeout: float,
    stdout_path: Path,
    stderr_path: Path,
    limite_saida: int = MAX_BYTES_SAIDA_OCR,
    arquivos_monitorados: tuple[Path, ...] = (),
    prazo: float | None = None,
) -> Any:
    """Executa Tesseract sem acumular saída e monitora processo e arquivos."""
    monitorados = _caminhos_unicos(
        tuple(arquivos_monitorados) + (stdout_path, stderr_path)
    )
    if prazo is None:
        prazo = monotonic() + timeout

    stdout = None
    stderr = None
    processo = None
    forcar_morte = False
    motivo_limite: str | None = None
    motivo_prazo = False
    try:
        stdout = _abrir_saida_regular(stdout_path, f"a saída stdout {stdout_path}")
        stderr = _abrir_saida_regular(stderr_path, f"a saída stderr {stderr_path}")
        processo = subprocess.Popen(argumentos, stdout=stdout, stderr=stderr)

        while True:
            try:
                retorno = processo.poll()
            except Exception:
                # O cleanup no finally ainda mata e aguarda o processo.
                raise
            tamanho = _tamanho_saida_fisico(
                monitorados,
                arquivos_abertos=(
                    (stdout_path, stdout),
                    (stderr_path, stderr),
                ),
            )
            agora = monotonic()

            if tamanho > limite_saida:
                motivo_limite = (
                    f"saída OCR combinada excede o limite de {limite_saida} bytes"
                )
            elif agora >= prazo:
                motivo_prazo = True

            if motivo_limite or motivo_prazo:
                # A exceção só é levantada depois de kill/wait no finally.
                if motivo_limite:
                    raise _OcrErro(motivo_limite)
                raise subprocess.TimeoutExpired(argumentos, timeout)

            if retorno is not None:
                concluido = subprocess.CompletedProcess(argumentos, retorno)
                # Mantém o fallback limitado para doubles que expõem stdout/
                # stderr em memória; o Popen real usa os arquivos acima.
                concluido.stdout = getattr(processo, "stdout", None)
                concluido.stderr = getattr(processo, "stderr", None)
                return concluido

            restante = prazo - monotonic()
            if restante <= 0:
                continue
            time.sleep(min(_INTERVALO_MONITORAMENTO_OCR, restante))
    except BaseException:
        # Uma saída inválida, limite, timeout ou erro inesperado ainda deixa o
        # processo em estado potencialmente vivo; force kill antes do wait.
        forcar_morte = True
        raise
    finally:
        try:
            if processo is not None:
                # Não importa se o loop saiu por exceção: poll/kill/wait ficam
                # centralizados aqui para não deixar subprocesso nem zumbi.
                _matar_e_aguardar(processo, forcar=forcar_morte)
        finally:
            try:
                if stdout is not None:
                    stdout.close()
            finally:
                if stderr is not None:
                    stderr.close()


def _valor_limitado(valor: object, limite: int, contexto: str) -> bytes:
    if valor is None:
        return b""
    if isinstance(valor, bytes):
        if len(valor) > limite:
            raise _OcrErro(f"{contexto} excede o limite de {limite} bytes")
        return valor
    if isinstance(valor, (bytearray, memoryview)):
        if len(valor) > limite:
            raise _OcrErro(f"{contexto} excede o limite de {limite} bytes")
        return bytes(valor)
    if isinstance(valor, str):
        # O retorno normal do subprocesso é None (o stream foi para arquivo).
        # Este caminho existe para mocks; rejeitar por caracteres antes de
        # codificar evita duplicar uma string arbitrariamente grande.
        if len(valor) > limite:
            raise _OcrErro(f"{contexto} excede o limite de {limite} bytes")
        dados = valor.encode("utf-8", errors="replace")
    else:
        representacao = str(valor)
        if len(representacao) > limite:
            raise _OcrErro(f"{contexto} excede o limite de {limite} bytes")
        dados = representacao.encode("utf-8", errors="replace")
    if len(dados) > limite:
        raise _OcrErro(f"{contexto} excede o limite de {limite} bytes")
    return dados


def _verificar_prazo(prazo: float | None, contexto: str) -> None:
    if prazo is not None and monotonic() >= prazo:
        raise _OcrErro(f"prazo esgotado durante {contexto}")


def _abrir_leitura_regular(caminho: Path, contexto: str) -> Any | None:
    status = _status_arquivo_regular(caminho, contexto)
    if status is None:
        return None

    try:
        arquivo = open(
            caminho,
            "rb",
            buffering=0,
            opener=_opener_seguro,
        )
    except FileNotFoundError:
        # O arquivo pode ser removido entre lstat e open; não se abre um
        # caminho substituído por outro tipo.
        return None
    except OSError as exc:
        raise _OcrErro(f"não foi possível abrir {contexto} com segurança: {exc}") from exc

    try:
        status_aberto = os.fstat(arquivo.fileno())
        if not stat.S_ISREG(status_aberto.st_mode):
            raise _OcrErro(f"{contexto} deve ser arquivo regular")
        return arquivo
    except Exception:
        arquivo.close()
        raise


def _ler_arquivo_limitado(
    caminho: Path,
    limite: int,
    contexto: str,
    *,
    prazo: float | None = None,
) -> bytes:
    """Lê em chunks, sem seguir links/FIFOs e respeitando o prazo."""
    if limite < 0:
        raise _OcrErro(f"{contexto} recebeu limite negativo")
    _verificar_prazo(prazo, f"a leitura de {contexto}")
    arquivo = _abrir_leitura_regular(caminho, contexto)
    if arquivo is None:
        _verificar_prazo(prazo, f"a leitura de {contexto}")
        return b""

    dados = bytearray()
    try:
        while True:
            _verificar_prazo(prazo, f"a leitura de {contexto}")
            try:
                bloco = arquivo.read(
                    min(_TAMANHO_CHUNK_OCR, limite - len(dados) + 1)
                )
            except BlockingIOError:
                if prazo is None:
                    raise _OcrErro(f"leitura não bloqueante indisponível para {contexto}")
                time.sleep(min(0.001, max(0.0, prazo - monotonic())))
                continue
            if bloco in (b"", None):
                if bloco is None:
                    raise _OcrErro(f"leitura inválida de {contexto}")
                break
            if not isinstance(bloco, (bytes, bytearray, memoryview)):
                raise _OcrErro(f"leitura inválida de {contexto}")
            dados.extend(bloco)
            if len(dados) > limite:
                raise _OcrErro(f"{contexto} excede o limite de {limite} bytes")
            try:
                tamanho_fisico = os.fstat(arquivo.fileno()).st_size
            except OSError as exc:
                raise _OcrErro(f"não foi possível verificar {contexto}: {exc}") from exc
            if tamanho_fisico > limite:
                raise _OcrErro(f"{contexto} excede o limite de {limite} bytes")
            _verificar_prazo(prazo, f"a leitura de {contexto}")
    finally:
        arquivo.close()
    return bytes(dados)


def _conteudo_limitado(
    caminho: Path,
    retorno: object,
    limite: int,
    contexto: str,
    *,
    prazo: float | None = None,
) -> str:
    """Lê a saída em chunks, com fallback limitado apenas para mocks."""
    try:
        dados = _ler_arquivo_limitado(
            caminho,
            limite,
            contexto,
            prazo=prazo,
        )
    except TypeError as exc:
        # Alguns testes/consumidores substituem o helper antigo de três
        # argumentos; ainda fazemos a checagem imediatamente depois da leitura.
        if "unexpected keyword argument 'prazo'" not in str(exc):
            raise
        dados = _ler_arquivo_limitado(caminho, limite, contexto)
    _verificar_prazo(prazo, f"a leitura de {contexto}")
    # Um mock pode devolver stdout em vez de escrever no descritor. Em uma
    # execução real stdout foi redirecionado para arquivo e normalmente é vazio;
    # ainda assim o fallback permanece limitado antes de ser decodificado.
    if not dados:
        dados = _valor_limitado(retorno, limite, contexto)
        _verificar_prazo(prazo, f"a leitura de {contexto}")
    conteudo = dados.decode("utf-8", errors="replace")
    _verificar_prazo(prazo, f"a decodificação de {contexto}")
    return conteudo


def _conteudo_limitado_com_prazo(
    caminho: Path,
    retorno: object,
    limite: int,
    contexto: str,
    prazo: float | None,
) -> str:
    """Chama o leitor com prazo, preservando doubles da assinatura antiga."""
    try:
        return _conteudo_limitado(
            caminho,
            retorno,
            limite,
            contexto,
            prazo=prazo,
        )
    except TypeError as exc:
        if "unexpected keyword argument 'prazo'" not in str(exc):
            raise
        return _conteudo_limitado(caminho, retorno, limite, contexto)


def _detalhe_subprocesso(
    resultado: Any,
    stderr_path: Path,
    limite: int,
    *,
    prazo: float | None = None,
) -> str:
    try:
        dados = _ler_arquivo_limitado(
            stderr_path,
            limite,
            "a saída de erro do OCR",
            prazo=prazo,
        )
    except _OcrErro as exc:
        if "excede o limite" in str(exc):
            return "saída de erro excede o limite"
        raise
    if not dados:
        try:
            dados = _valor_limitado(
                getattr(resultado, "stderr", None),
                limite,
                "a saída de erro do OCR",
            )
        except _OcrErro as exc:
            if "excede o limite" in str(exc):
                return "saída de erro excede o limite"
            raise
    _verificar_prazo(prazo, "a leitura da saída de erro do OCR")
    return dados.decode("utf-8", errors="replace").strip()[:4096]


def _normalizar_texto(texto: str) -> str:
    """Normaliza espaços para que a comparação não premie padding."""
    return " ".join(texto.split())


def _metricas_texto(texto: str) -> tuple[int, int]:
    normalizado = _normalizar_texto(texto)
    return len(normalizado), len(normalizado.split())


def _qualidade_texto(
    texto: str,
    confianca_media: float | None = None,
    *,
    palavras: int | None = None,
) -> _QualidadeTexto:
    limpo = _normalizar_texto(texto)
    caracteres = len(limpo)
    quantidade_palavras = len(limpo.split()) if palavras is None else palavras
    motivos: list[str] = []
    if caracteres < MIN_CARACTERES_POR_PAGINA:
        motivos.append(
            f"{caracteres} caracteres (mínimo {MIN_CARACTERES_POR_PAGINA})"
        )
    if quantidade_palavras < MIN_PALAVRAS_POR_PAGINA:
        motivos.append(
            f"{quantidade_palavras} palavras (mínimo {MIN_PALAVRAS_POR_PAGINA})"
        )
    try:
        confianca_numerica = (
            None if confianca_media is None else float(confianca_media)
        )
    except (TypeError, ValueError):
        confianca_numerica = None
    if (
        confianca_numerica is None
        or not math.isfinite(confianca_numerica)
        or confianca_numerica < 0
        or confianca_numerica > 100
    ):
        motivos.append("confiança ausente ou inválida")
    elif confianca_numerica < MIN_CONFIANCA_OCR:
        motivos.append(
            f"confiança média {confianca_numerica:.2f} (mínimo {MIN_CONFIANCA_OCR:g})"
        )
    return _QualidadeTexto(
        caracteres=caracteres,
        palavras=quantidade_palavras,
        confianca_media=confianca_numerica,
        passou=not motivos,
        motivos=tuple(motivos),
    )


def _texto_insuficiente(texto: str) -> bool:
    """Decide se a extração direta é curta ou estruturalmente corrompida."""
    normalizado = _normalizar_texto(texto)
    if len(normalizado) < MIN_CARACTERES_POR_PAGINA:
        return True
    controles = sum(
        unicodedata.category(caractere).startswith("C")
        for caractere in normalizado
    )
    return (
        controles >= MIN_CARACTERES_CONTROLE_CORROMPIDOS
        and controles / len(normalizado) > MAX_FRACAO_CARACTERES_CONTROLE
    )


def _dimensoes_pagina_pontos(pagina: Any) -> tuple[float, float] | None:
    try:
        retangulo = getattr(pagina, "rect", None)
        if callable(retangulo):
            retangulo = retangulo()
        largura = float(getattr(retangulo, "width"))
        altura = float(getattr(retangulo, "height"))
    except (AttributeError, TypeError, ValueError):
        return None
    if not math.isfinite(largura) or not math.isfinite(altura):
        return None
    if largura <= 0 or altura <= 0:
        return None
    return largura, altura


def _matriz_raster(
    pagina: Any,
    pymupdf: Any,
    *,
    dpi: int,
    max_megapixels: float,
) -> Any:
    escala = dpi / 72.0
    dimensoes = _dimensoes_pagina_pontos(pagina)
    max_pixels = max(1, math.floor(max_megapixels * 1_000_000))
    if dimensoes is not None:
        largura, altura = dimensoes
        largura_pixels = max(1, math.ceil(largura * escala))
        altura_pixels = max(1, math.ceil(altura * escala))
        pixels_previstos = largura_pixels * altura_pixels
        if pixels_previstos > max_pixels:
            escala *= math.sqrt(max_pixels / pixels_previstos)
            # As dimensões finais do Pixmap são arredondadas de novo pelo
            # MuPDF. Recalcular algumas vezes deixa margem para esse arredondamento
            # e evita criar um raster ligeiramente acima do teto.
            for _ in range(3):
                largura_final = max(1, math.ceil(largura * escala))
                altura_final = max(1, math.ceil(altura * escala))
                pixels_finais = largura_final * altura_final
                if pixels_finais <= max_pixels:
                    break
                escala *= math.sqrt(max_pixels / pixels_finais) * 0.999999
    return pymupdf.Matrix(escala, escala)


def _gravar_pixmap(
    pixmap: Any,
    caminho: Path,
    *,
    max_megapixels: float,
) -> None:
    max_pixels = max(1, math.floor(max_megapixels * 1_000_000))
    try:
        largura = int(getattr(pixmap, "width"))
        altura = int(getattr(pixmap, "height"))
    except (AttributeError, TypeError, ValueError):
        largura = altura = 0
    if largura > 0 and altura > 0 and largura * altura > max_pixels:
        raise _OcrErro(
            f"raster da página excede o limite de {max_megapixels:g} megapixels"
        )

    contexto_imagem = "a imagem rasterizada da página"
    # A imagem também é uma saída temporária consumida pelo filho: não
    # permitimos que uma troca por link ou especial faça o Tesseract bloquear.
    _status_arquivo_regular(caminho, contexto_imagem)
    salvar = getattr(pixmap, "save", None)
    try:
        if callable(salvar):
            salvar(str(caminho))
        if not caminho.exists():
            raise OSError("Pixmap.save não produziu o arquivo de imagem")
        _status_arquivo_regular(caminho, contexto_imagem, ausente_ok=False)
    except _OcrErro:
        raise
    except Exception as erro_salvar:
        # O fallback facilita mocks determinísticos e versões de PyMuPDF que
        # expõem apenas tobytes; a imagem continua indo para o temporário.
        try:
            _status_arquivo_regular(caminho, contexto_imagem)
            tobytes = getattr(pixmap, "tobytes")
            try:
                dados = tobytes("png")
            except TypeError:
                dados = tobytes()
            if not isinstance(dados, (bytes, bytearray, memoryview)):
                raise TypeError("Pixmap.tobytes não retornou bytes")
            arquivo_imagem = _abrir_saida_regular(caminho, contexto_imagem)
            try:
                arquivo_imagem.write(bytes(dados))
            finally:
                arquivo_imagem.close()
            _status_arquivo_regular(caminho, contexto_imagem, ausente_ok=False)
        except Exception as exc:
            raise _OcrErro(
                f"falha ao rasterizar a página para OCR: "
                f"{type(erro_salvar).__name__}: {erro_salvar}; "
                f"fallback {type(exc).__name__}: {exc}"
            ) from exc


def _parsear_tsv(
    conteudo: str,
    numero: int,
    *,
    prazo: float | None = None,
) -> _ResultadoOcrPagina:
    """Parseia TSV estrito; qualquer ``csv.Error`` invalida só esta página."""
    try:
        leitor = csv.DictReader(
            io.StringIO(conteudo),
            delimiter="\t",
            strict=True,
        )
        campos = leitor.fieldnames or []
        campos_normalizados = {
            str(campo).lstrip("\ufeff").strip().lower(): campo
            for campo in campos
            if campo is not None
        }
        campo_texto = campos_normalizados.get("text")
        if campo_texto is None:
            raise _OcrErro(
                f"Tesseract não produziu a coluna text no TSV da página {numero}"
            )
        campo_nivel = campos_normalizados.get("level")
        campo_confianca = campos_normalizados.get("conf")

        linhas: list[tuple[str, object]] = []
        linhas_nivel_palavra: list[tuple[str, object]] = []
        for linha in leitor:
            _verificar_prazo(prazo, f"o parsing TSV da página {numero}")
            if not isinstance(linha, dict):
                continue
            if linha.get(None):
                raise _OcrErro(f"TSV inválido na página {numero}: colunas extras")
            valor = linha.get(campo_texto)
            texto = "" if valor is None else str(valor).strip()
            if not texto:
                continue
            confianca: object = linha.get(campo_confianca) if campo_confianca else None
            registro = (texto, confianca)
            linhas.append(registro)
            nivel = str(linha.get(campo_nivel, "")).strip() if campo_nivel else ""
            if nivel == "5":
                linhas_nivel_palavra.append(registro)
    except (csv.Error, TypeError, ValueError) as exc:
        raise _OcrErro(
            f"TSV inválido na página {numero}: {type(exc).__name__}: {exc}"
        ) from exc

    # Tesseract emite palavras no level 5. O fallback para todas as linhas
    # torna o parser tolerante a TSVs sintéticos usados pelos testes.
    selecionadas = linhas_nivel_palavra or linhas
    _verificar_prazo(prazo, f"o parsing TSV da página {numero}")
    texto = " ".join(valor for valor, _ in selecionadas).strip()
    palavras = _metricas_texto(texto)[1]
    confiancas: list[float] = []
    confianca_invalida = False
    for _, valor in selecionadas:
        _verificar_prazo(prazo, f"o parsing TSV da página {numero}")
        try:
            numero_confianca = float(str(valor).strip())
        except (TypeError, ValueError):
            confianca_invalida = True
            continue
        # -1 é o marcador do Tesseract para níveis sem confiança. Valores
        # não finitos também não podem compor uma média aceitável.
        if (
            not math.isfinite(numero_confianca)
            or numero_confianca < 0
            or numero_confianca > 100
        ):
            confianca_invalida = True
            continue
        confiancas.append(numero_confianca)
    _verificar_prazo(prazo, f"o parsing TSV da página {numero}")
    confianca_media = (
        sum(confiancas) / len(confiancas)
        if confiancas and not confianca_invalida and len(confiancas) == len(selecionadas)
        else None
    )
    return _ResultadoOcrPagina(texto, palavras, confianca_media)


def _prazo_restante(prazo: float | None) -> float | None:
    if prazo is None:
        return None
    return prazo - monotonic()


def _verificar_prazos_pagina(
    numero: int,
    prazo_total: float | None,
    prazo_pagina: float | None,
    etapa: str,
) -> None:
    restante_total = _prazo_restante(prazo_total)
    if restante_total is not None and restante_total <= 0:
        raise _OcrErro(
            f"tempo total de OCR esgotado durante {etapa} da página {numero}"
        )
    restante_pagina = _prazo_restante(prazo_pagina)
    if restante_pagina is not None and restante_pagina <= 0:
        raise _OcrErro(
            f"tempo de OCR da página {numero} esgotado durante {etapa}"
        )


def _parsear_tsv_com_prazo(
    conteudo: str,
    numero: int,
    prazo: float | None,
) -> _ResultadoOcrPagina:
    """Mantém compatibilidade com parsers de teste que aceitam dois args."""
    try:
        return _parsear_tsv(conteudo, numero, prazo=prazo)
    except TypeError as exc:
        if "unexpected keyword argument 'prazo'" not in str(exc):
            raise
        return _parsear_tsv(conteudo, numero)


def _ocr_pagina(
    pagina: Any,
    numero: int,
    executavel: str,
    idioma: str,
    pymupdf: Any,
    *,
    diretorio: Path | None = None,
    dpi: int = DPI_OCR,
    max_megapixels: float = MAX_MEGAPIXELS_OCR,
    limite_saida: int = MAX_BYTES_SAIDA_OCR,
    timeout: float = MAX_TEMPO_OCR_PAGINA,
    prazo_total: float | None = None,
    prazo_pagina: float | None = None,
) -> _ResultadoOcrPagina:
    """Rasteriza uma página e obtém seu TSV em um temporário isolado."""
    if diretorio is None:
        with tempfile.TemporaryDirectory(prefix="licita-corpus-ocr-") as temporario:
            return _ocr_pagina(
                pagina,
                numero,
                executavel,
                idioma,
                pymupdf,
                diretorio=Path(temporario),
                dpi=dpi,
                max_megapixels=max_megapixels,
                limite_saida=limite_saida,
                timeout=timeout,
                prazo_total=prazo_total,
                prazo_pagina=prazo_pagina,
            )

    _verificar_prazos_pagina(numero, prazo_total, prazo_pagina, "a preparação")
    try:
        matriz = _matriz_raster(
            pagina,
            pymupdf,
            dpi=dpi,
            max_megapixels=max_megapixels,
        )
        pixmap = pagina.get_pixmap(matrix=matriz, alpha=False)
        imagem_path = diretorio / f"pagina-{numero}.png"
        _gravar_pixmap(pixmap, imagem_path, max_megapixels=max_megapixels)
    except _OcrErro:
        raise
    except Exception as exc:
        raise _OcrErro(
            f"falha ao rasterizar a página {numero} com PyMuPDF: "
            f"{type(exc).__name__}: {exc}"
        ) from exc
    _verificar_prazos_pagina(numero, prazo_total, prazo_pagina, "a rasterização")

    # O processo, a leitura e o parser compartilham o menor dos prazos.
    agora_execucao = monotonic()
    prazos = [
        prazo
        for prazo in (prazo_total, prazo_pagina)
        if prazo is not None
    ]
    prazos.append(agora_execucao + timeout)
    prazo_execucao = min(prazos)
    timeout_processo = prazo_execucao - agora_execucao
    if timeout_processo <= 0:
        _verificar_prazos_pagina(numero, prazo_total, prazo_pagina, "a execução")
        raise _OcrErro(f"tempo de OCR da página {numero} esgotado antes da execução")

    base_saida = diretorio / f"pagina-{numero}"
    stdout_path = diretorio / f"pagina-{numero}.stdout"
    stderr_path = diretorio / f"pagina-{numero}.stderr"
    caminhos_tsv = _caminhos_unicos(
        (
            base_saida.with_suffix(".tsv"),
            Path(f"{base_saida}.tsv"),
            base_saida,
        )
    )
    # O arquivo principal é criado antes do filho: isso impede que uma saída
    # ausente seja confundida com FIFO/link e também entra no monitor físico.
    for caminho_tsv in caminhos_tsv:
        _status_arquivo_regular(
            caminho_tsv,
            f"a saída TSV OCR da página {numero} ({caminho_tsv.name})",
        )
    tsv_principal = caminhos_tsv[0]
    _criar_saida_regular(
        tsv_principal,
        f"a saída TSV OCR da página {numero}",
    )

    argumentos = [
        executavel,
        str(imagem_path),
        str(base_saida),
        "-l",
        idioma,
        "tsv",
    ]
    try:
        resultado = _rodar_subprocesso(
            argumentos,
            timeout=timeout_processo,
            stdout_path=stdout_path,
            stderr_path=stderr_path,
            limite_saida=limite_saida,
            arquivos_monitorados=caminhos_tsv,
            prazo=prazo_execucao,
        )
    except subprocess.TimeoutExpired as exc:
        raise _OcrErro(
            f"TimeoutExpired: tempo esgotado no Tesseract para a página {numero} "
            f"(limite efetivo {timeout_processo:.6g}s)"
        ) from exc
    except FileNotFoundError as exc:
        raise _OcrErro(
            f"executável do Tesseract não encontrado ao processar a página {numero}: "
            f"{executavel!r}"
        ) from exc
    except OSError as exc:
        raise _OcrErro(
            f"não foi possível executar o Tesseract na página {numero}: {exc}"
        ) from exc
    except Exception as exc:
        raise _OcrErro(
            f"erro ao executar o Tesseract na página {numero}: "
            f"{type(exc).__name__}: {exc}"
        ) from exc

    # Revalida depois do wait: um fake ou um write final pode ter crescido fora
    # do último intervalo de poll.
    tamanho_final = _tamanho_saida_fisico(
        caminhos_tsv + (stderr_path, stdout_path)
    )
    if tamanho_final > limite_saida:
        raise _OcrErro(
            f"saída OCR combinada excede o limite de {limite_saida} bytes"
        )
    _verificar_prazo(prazo_execucao, f"a conclusão da página {numero}")
    if getattr(resultado, "returncode", 1) != 0:
        detalhe = _detalhe_subprocesso(
            resultado,
            stderr_path,
            limite_saida,
            prazo=prazo_execucao,
        )
        sufixo = f": {detalhe}" if detalhe else ""
        raise _OcrErro(
            f"Tesseract retornou código não-zero na página {numero}{sufixo}"
        )

    # O resultado só é utilizável se o deadline sobreviver ao processo, à
    # leitura e ao parser; nenhum estágio pode estender o orçamento.
    _verificar_prazos_pagina(numero, prazo_total, prazo_pagina, "a execução")
    status_tsv = {
        caminho: _status_arquivo_regular(
            caminho,
            f"a saída TSV OCR da página {numero} ({caminho.name})",
        )
        for caminho in caminhos_tsv
    }
    tsv_path = next(
        (
            caminho
            for caminho in caminhos_tsv
            if status_tsv[caminho] is not None and status_tsv[caminho].st_size > 0
        ),
        None,
    )
    if tsv_path is None:
        tsv_path = next(
            (caminho for caminho in caminhos_tsv if status_tsv[caminho] is not None),
            None,
        )

    contexto_tsv = f"a saída TSV OCR da página {numero}"
    if tsv_path is not None:
        conteudo = _conteudo_limitado_com_prazo(
            tsv_path,
            None,
            limite_saida,
            contexto_tsv,
            prazo_execucao,
        )
        if not conteudo:
            conteudo = _conteudo_limitado_com_prazo(
                stdout_path,
                getattr(resultado, "stdout", None),
                limite_saida,
                contexto_tsv,
                prazo_execucao,
            )
    else:
        conteudo = _conteudo_limitado_com_prazo(
            stdout_path,
            getattr(resultado, "stdout", None),
            limite_saida,
            contexto_tsv,
            prazo_execucao,
        )
    _verificar_prazos_pagina(numero, prazo_total, prazo_pagina, "a leitura")
    if not conteudo:
        raise _OcrErro(f"Tesseract não produziu saída TSV para a página {numero}")

    try:
        resultado_tsv = _parsear_tsv_com_prazo(
            conteudo,
            numero,
            prazo_execucao,
        )
    except _OcrErro:
        raise
    except (csv.Error, TypeError, ValueError) as exc:
        raise _OcrErro(
            f"TSV inválido na página {numero}: {type(exc).__name__}: {exc}"
        ) from exc
    _verificar_prazos_pagina(numero, prazo_total, prazo_pagina, "o parsing")
    _verificar_prazo(prazo_execucao, f"o parsing da página {numero}")
    return resultado_tsv


def _avaliar_paginas(
    partes: list[str],
    originais: list[str],
    paginas_com_ocr: set[int],
    confiancas: Mapping[int, float | None] | None = None,
    erros_por_pagina: Mapping[int, str] | None = None,
) -> tuple[AvaliacaoPagina, ...]:
    confiancas = confiancas or {}
    erros_por_pagina = erros_por_pagina or {}
    return tuple(
        AvaliacaoPagina(
            pagina=numero,
            caracteres=_metricas_texto(texto)[0],
            precisa_ocr=_texto_insuficiente(texto),
            ocr_usado=numero in paginas_com_ocr,
            caracteres_originais=_metricas_texto(original)[0],
            texto=texto,
            texto_original=original,
            palavras=_metricas_texto(texto)[1],
            palavras_originais=_metricas_texto(original)[1],
            confianca_media=confiancas.get(numero),
            erro_ocr=erros_por_pagina.get(numero),
        )
        for numero, (texto, original) in enumerate(zip(partes, originais), start=1)
    )


def _executar_ocr_paginas_com_objetos(
    partes: list[str],
    originais: list[str],
    objetos_paginas: list[Any],
    paginas_que_precisam: set[int],
    *,
    executavel: str,
    idioma: str,
    componentes: tuple[str, ...],
    pymupdf: Any,
    inicio: float,
    limite_paginas: int,
    limite_total: float,
    limite_pagina: float,
    dpi: int,
    max_megapixels: float,
    limite_saida: int,
) -> _ResultadoOcrDocumento:
    """Executa OCR mantendo o temporário de cada página isolado."""
    prazo_total = inicio + limite_total
    paginas_tentadas: set[int] = set()
    paginas_aplicadas: set[int] = set()
    confiancas: dict[int, float | None] = {}
    erros: list[str] = []
    erros_por_pagina: dict[int, str] = {}
    palavras_ocr: dict[int, int] = {}

    candidatas = sorted(paginas_que_precisam)
    if len(candidatas) > limite_paginas:
        excedentes = len(candidatas) - limite_paginas
        erros.append(
            f"limite de {limite_paginas} páginas OCR atingido; "
            f"{excedentes} página(s) não processada(s)"
        )
        candidatas = candidatas[:limite_paginas]
    if not candidatas:
        return _ResultadoOcrDocumento(
            frozenset(), frozenset(), confiancas, tuple(erros)
        )

    restante = _prazo_restante(prazo_total)
    if restante is not None and restante <= 0:
        erros.append("tempo total de OCR esgotado antes de iniciar o Tesseract")
        return _ResultadoOcrDocumento(
            frozenset(), frozenset(), confiancas, tuple(erros)
        )

    # A consulta de idiomas não deve manter arquivos junto com as páginas.
    try:
        with tempfile.TemporaryDirectory(prefix="licita-corpus-ocr-langs-") as temporario:
            _validar_tesseract(
                executavel,
                idioma,
                componentes,
                timeout=restante if restante is not None else limite_total,
                diretorio=Path(temporario),
                limite_saida=limite_saida,
                prazo=prazo_total,
            )
    except _OcrErro as exc:
        erros.append(str(exc))
        return _ResultadoOcrDocumento(
            frozenset(paginas_aplicadas),
            frozenset(paginas_tentadas),
            confiancas,
            tuple(erros),
        )

    try:
        _verificar_prazo(prazo_total, "a validação de idiomas")
    except _OcrErro as exc:
        erros.append(str(exc))
        return _ResultadoOcrDocumento(
            frozenset(paginas_aplicadas),
            frozenset(paginas_tentadas),
            confiancas,
            tuple(erros),
        )

    for numero in candidatas:
        restante = _prazo_restante(prazo_total)
        if restante is not None and restante <= 0:
            erros.append(f"tempo total de OCR esgotado antes da página {numero}")
            break
        inicio_pagina = monotonic()
        restante_pagina = prazo_total - inicio_pagina
        if restante_pagina <= 0:
            erros.append(f"tempo total de OCR esgotado antes da página {numero}")
            break
        prazo_pagina = inicio_pagina + min(limite_pagina, restante_pagina)
        paginas_tentadas.add(numero)
        try:
            # Cada página tem diretório próprio e o contexto termina antes de
            # a próxima página começar, removendo imagem, TSV e stderr.
            with tempfile.TemporaryDirectory(
                prefix=f"licita-corpus-ocr-p{numero}-"
            ) as temporario_pagina:
                resultado = _ocr_pagina(
                    objetos_paginas[numero - 1],
                    numero,
                    executavel,
                    idioma,
                    pymupdf,
                    diretorio=Path(temporario_pagina),
                    dpi=dpi,
                    max_megapixels=max_megapixels,
                    limite_saida=limite_saida,
                    timeout=min(limite_pagina, restante_pagina),
                    prazo_total=prazo_total,
                    prazo_pagina=prazo_pagina,
                )
        except _OcrErro as exc:
            mensagem = str(exc)
            erros.append(mensagem)
            erros_por_pagina[numero] = mensagem
            continue
        except Exception as exc:
            mensagem = (
                f"erro inesperado no OCR da página {numero}: "
                f"{type(exc).__name__}: {exc}"
            )
            erros.append(mensagem)
            erros_por_pagina[numero] = mensagem
            continue

        confiancas[numero] = resultado.confianca_media
        palavras_ocr[numero] = resultado.palavras
        try:
            _verificar_prazos_pagina(
                numero,
                prazo_total,
                prazo_pagina,
                "a aceitação",
            )
        except _OcrErro as exc:
            mensagem = str(exc)
            erros.append(mensagem)
            erros_por_pagina[numero] = mensagem
            continue

        qualidade = _qualidade_texto(
            resultado.texto,
            resultado.confianca_media,
            palavras=resultado.palavras,
        )
        motivos = list(qualidade.motivos)
        caracteres_originais, palavras_originais = _metricas_texto(originais[numero - 1])
        if qualidade.caracteres < caracteres_originais:
            motivos.append(
                f"{qualidade.caracteres} caracteres não melhora os "
                f"{caracteres_originais} caracteres originais normalizados"
            )
        if qualidade.palavras < palavras_originais:
            motivos.append(
                f"{qualidade.palavras} palavras reduz as "
                f"{palavras_originais} palavras originais normalizadas"
            )
        if motivos:
            mensagem = f"OCR rejeitado na página {numero}: " + "; ".join(motivos)
            erros.append(mensagem)
            erros_por_pagina[numero] = mensagem
            # A página permanece exatamente com seu texto original.
            continue

        _verificar_prazos_pagina(
            numero,
            prazo_total,
            prazo_pagina,
            "a substituição",
        )
        # Só esta página é substituída; o PDF fonte e as demais páginas ficam
        # intactos.
        partes[numero - 1] = resultado.texto
        paginas_aplicadas.add(numero)

    return _ResultadoOcrDocumento(
        frozenset(paginas_aplicadas),
        frozenset(paginas_tentadas),
        confiancas,
        tuple(erros),
        palavras_ocr,
        erros_por_pagina,
    )


def _media_confiancas(
    confiancas: Mapping[int, float | None],
    paginas: Mapping[int, int] | None = None,
) -> float | None:
    valores: list[tuple[float, int]] = []
    for pagina, confianca in confiancas.items():
        if confianca is None:
            continue
        try:
            confianca_numerica = float(confianca)
        except (TypeError, ValueError):
            continue
        if (
            not math.isfinite(confianca_numerica)
            or confianca_numerica < 0
            or confianca_numerica > 100
        ):
            continue
        peso = (paginas or {}).get(pagina, 1)
        valores.append((confianca_numerica, max(1, peso)))
    if not valores:
        return None
    total_peso = sum(peso for _, peso in valores)
    return sum(valor * peso for valor, peso in valores) / total_peso


def _primeiro_configurado(*valores: object, padrao: object) -> object:
    for valor in valores:
        if valor is not None:
            return valor
    return padrao


def _megapixels_de_pixels(valor: object) -> float | None:
    if valor is None:
        return None
    if isinstance(valor, bool) or not isinstance(valor, (int, float)):
        return valor  # a validação principal produz a mensagem adequada
    return float(valor) / 1_000_000


def _validar_configuracao(
    *,
    max_paginas_ocr: object,
    tempo_total_ocr: object,
    tempo_pagina_ocr: object,
    dpi_ocr: object,
    max_megapixels_ocr: object,
    max_saida_ocr: object,
) -> tuple[int, float, float, int, float, int]:
    if isinstance(max_paginas_ocr, bool) or not isinstance(max_paginas_ocr, int):
        raise ValueError("max_paginas_ocr deve ser um inteiro")
    if max_paginas_ocr < 0:
        raise ValueError("max_paginas_ocr não pode ser negativo")
    if isinstance(tempo_total_ocr, bool) or not isinstance(tempo_total_ocr, (int, float)):
        raise ValueError("tempo_total_ocr deve ser numérico")
    if not math.isfinite(float(tempo_total_ocr)) or float(tempo_total_ocr) < 0:
        raise ValueError("tempo_total_ocr deve ser finito e não negativo")
    if isinstance(tempo_pagina_ocr, bool) or not isinstance(tempo_pagina_ocr, (int, float)):
        raise ValueError("tempo_pagina_ocr deve ser numérico")
    if not math.isfinite(float(tempo_pagina_ocr)) or float(tempo_pagina_ocr) < 0:
        raise ValueError("tempo_pagina_ocr deve ser finito e não negativo")
    if isinstance(dpi_ocr, bool) or not isinstance(dpi_ocr, int) or dpi_ocr <= 0:
        raise ValueError("dpi_ocr deve ser um inteiro positivo")
    if isinstance(max_megapixels_ocr, bool) or not isinstance(max_megapixels_ocr, (int, float)):
        raise ValueError("max_megapixels_ocr deve ser numérico")
    if not math.isfinite(float(max_megapixels_ocr)) or float(max_megapixels_ocr) <= 0:
        raise ValueError("max_megapixels_ocr deve ser positivo")
    if isinstance(max_saida_ocr, bool) or not isinstance(max_saida_ocr, int) or max_saida_ocr <= 0:
        raise ValueError("max_saida_ocr deve ser um inteiro positivo")
    return (
        max_paginas_ocr,
        float(tempo_total_ocr),
        float(tempo_pagina_ocr),
        dpi_ocr,
        float(max_megapixels_ocr),
        max_saida_ocr,
    )


def _verificar_pdf(
    caminho: Path,
    *,
    usar_ocr: bool = False,
    idioma: str = "por",
    executavel_tesseract: str = "tesseract",
    max_paginas_ocr: int = MAX_PAGINAS_OCR,
    tempo_total_ocr: float = MAX_TEMPO_OCR_DOCUMENTO,
    tempo_pagina_ocr: float = MAX_TEMPO_OCR_PAGINA,
    dpi_ocr: int = DPI_OCR,
    max_megapixels_ocr: float = MAX_MEGAPIXELS_OCR,
    max_saida_ocr: int = MAX_BYTES_SAIDA_OCR,
) -> Verificacao:
    import pymupdf

    idioma_normalizado: str | None = None
    erros: list[str] = []
    paginas_com_ocr: set[int] = set()
    paginas_tentadas: set[int] = set()
    confiancas: dict[int, float | None] = {}
    palavras_ocr: dict[int, int] = {}
    erros_por_pagina: dict[int, str] = {}
    inicio_ocr = monotonic() if usar_ocr else None
    if usar_ocr:
        try:
            idioma_normalizado, componentes = _idioma_ocr(idioma)
        except _OcrErro as exc:
            erros.append(str(exc))
            componentes = ()
    else:
        componentes = ()

    try:
        with pymupdf.open(caminho) as documento:
            paginas = documento.page_count
            objetos_paginas = list(documento)
            originais = [pagina.get_text() for pagina in objetos_paginas]
            partes = list(originais)

            paginas_que_precisam = {
                numero
                for numero, texto in enumerate(originais, start=1)
                if _texto_insuficiente(texto)
            }
            if usar_ocr and not erros and paginas_que_precisam:
                try:
                    resultado_ocr = _executar_ocr_paginas_com_objetos(
                        partes,
                        originais,
                        objetos_paginas,
                        paginas_que_precisam,
                        executavel=executavel_tesseract,
                        idioma=idioma_normalizado or idioma,
                        componentes=componentes,
                        pymupdf=pymupdf,
                        inicio=inicio_ocr if inicio_ocr is not None else monotonic(),
                        limite_paginas=max_paginas_ocr,
                        limite_total=tempo_total_ocr,
                        limite_pagina=tempo_pagina_ocr,
                        dpi=dpi_ocr,
                        max_megapixels=max_megapixels_ocr,
                        limite_saida=max_saida_ocr,
                    )
                    paginas_com_ocr.update(resultado_ocr.paginas_aplicadas)
                    paginas_tentadas.update(resultado_ocr.paginas_tentadas)
                    confiancas.update(resultado_ocr.confiancas)
                    palavras_ocr.update(resultado_ocr.palavras)
                    erros_por_pagina.update(resultado_ocr.erros_por_pagina)
                    erros.extend(resultado_ocr.erros)
                except Exception as exc:
                    # PDF válido não vira erro de abertura por uma falha do
                    # fallback. O texto original permanece em ``partes``.
                    erros.append(
                        f"erro inesperado no OCR: {type(exc).__name__}: {exc}"
                    )
    except Exception as exc:
        mensagens = list(erros)
        mensagens.append(f"{type(exc).__name__}: {exc}")
        return Verificacao(
            caminho,
            abriu=False,
            erro="; ".join(mensagens),
            ocr_solicitado=usar_ocr,
            ocr_motor="tesseract" if usar_ocr else None,
            ocr_idioma=idioma_normalizado if usar_ocr else None,
            ocr_erros=tuple(erros),
        )

    avaliadas = _avaliar_paginas(
        partes,
        originais,
        paginas_com_ocr,
        confiancas,
        erros_por_pagina,
    )
    texto = "\n".join(partes)
    texto_original = "\n".join(originais)
    precisa_ocr = not avaliadas or any(pagina.precisa_ocr for pagina in avaliadas)
    return Verificacao(
        caminho,
        abriu=True,
        paginas=paginas,
        caracteres=len(texto.strip()),
        precisa_ocr=precisa_ocr,
        erro="; ".join(erros) or None,
        texto=texto,
        ocr_solicitado=usar_ocr,
        ocr_usado=bool(paginas_com_ocr),
        ocr_motor="tesseract" if usar_ocr else None,
        ocr_idioma=idioma_normalizado if usar_ocr else None,
        paginas_ocr=tuple(sorted(paginas_com_ocr)),
        paginas_avaliadas=avaliadas,
        texto_original=texto_original,
        ocr_confianca_media=_media_confiancas(confiancas, palavras_ocr),
        ocr_erros=tuple(erros),
        paginas_ocr_tentadas=tuple(sorted(paginas_tentadas)),
    )


def _verificar_docx(caminho: Path) -> Verificacao:
    import docx

    try:
        documento = docx.Document(str(caminho))
        partes = [p.text for p in documento.paragraphs]
        for tabela in documento.tables:
            for linha in tabela.rows:
                partes.extend(celula.text for celula in linha.cells)
    except Exception as exc:
        return Verificacao(caminho, abriu=False, erro=f"{type(exc).__name__}: {exc}")
    texto = "\n".join(partes)
    return Verificacao(
        caminho,
        abriu=True,
        paginas=None,
        caracteres=len(texto.strip()),
        texto=texto,
        texto_original=texto,
    )


def verificar(
    caminho: Path,
    *,
    ocr: bool = False,
    idioma: str = "por",
    executavel_tesseract: str = "tesseract",
    usar_ocr: bool | None = None,
    idioma_ocr: str | None = None,
    max_paginas_ocr: int | None = None,
    tempo_total_ocr: float | None = None,
    tempo_pagina_ocr: float | None = None,
    dpi_ocr: int | None = None,
    max_megapixels_ocr: float | None = None,
    max_saida_ocr: int | None = None,
    # Nomes alternativos mantêm a configuração fácil de descobrir em código
    # que usa convenções em inglês ou chama o limite diretamente de timeout.
    max_pages: int | None = None,
    max_pages_ocr: int | None = None,
    max_ocr_pages: int | None = None,
    ocr_max_pages: int | None = None,
    timeout_total: float | None = None,
    timeout_total_ocr: float | None = None,
    ocr_timeout_total: float | None = None,
    ocr_total_timeout: float | None = None,
    timeout_page: float | None = None,
    timeout_page_ocr: float | None = None,
    ocr_timeout_page: float | None = None,
    ocr_page_timeout: float | None = None,
    dpi: int | None = None,
    ocr_dpi: int | None = None,
    max_megapixels: float | None = None,
    ocr_max_megapixels: float | None = None,
    max_ocr_megapixels: float | None = None,
    max_pixels: int | None = None,
    max_raster_pixels: int | None = None,
    ocr_max_pixels: int | None = None,
    max_output_bytes: int | None = None,
    max_output_bytes_ocr: int | None = None,
    ocr_max_output_bytes: int | None = None,
) -> Verificacao:
    """Abre um documento e, opcionalmente, aplica OCR às páginas insuficientes.

    ``verificar(caminho)`` mantém o fluxo compatível e não executa OCR. Com
    ``ocr=True``, somente páginas abaixo dos critérios mínimos são rasterizadas.
    Cada resultado TSV precisa passar caracteres, palavras e, quando o TSV
    trouxer confiança, o limiar de confiança; caso contrário, o texto original
    daquela página é preservado.
    """
    if usar_ocr is not None:
        ocr = usar_ocr
    if idioma_ocr is not None:
        idioma = idioma_ocr

    caminho = Path(caminho)
    extensao = caminho.suffix.lower().lstrip(".")
    if extensao == "pdf":
        configuracao = _validar_configuracao(
            max_paginas_ocr=_primeiro_configurado(
                max_paginas_ocr,
                max_pages_ocr,
                max_ocr_pages,
                ocr_max_pages,
                max_pages,
                padrao=MAX_PAGINAS_OCR,
            ),
            tempo_total_ocr=_primeiro_configurado(
                tempo_total_ocr,
                timeout_total_ocr,
                ocr_timeout_total,
                ocr_total_timeout,
                timeout_total,
                padrao=MAX_TEMPO_OCR_DOCUMENTO,
            ),
            tempo_pagina_ocr=_primeiro_configurado(
                tempo_pagina_ocr,
                timeout_page_ocr,
                ocr_timeout_page,
                ocr_page_timeout,
                timeout_page,
                padrao=MAX_TEMPO_OCR_PAGINA,
            ),
            dpi_ocr=_primeiro_configurado(dpi_ocr, ocr_dpi, dpi, padrao=DPI_OCR),
            max_megapixels_ocr=_primeiro_configurado(
                max_megapixels_ocr,
                ocr_max_megapixels,
                max_ocr_megapixels,
                _megapixels_de_pixels(
                    _primeiro_configurado(
                        ocr_max_pixels,
                        max_raster_pixels,
                        max_pixels,
                        padrao=None,
                    )
                ),
                max_megapixels,
                padrao=MAX_MEGAPIXELS_OCR,
            ),
            max_saida_ocr=_primeiro_configurado(
                max_saida_ocr,
                max_output_bytes_ocr,
                ocr_max_output_bytes,
                max_output_bytes,
                padrao=MAX_BYTES_SAIDA_OCR,
            ),
        )
        return _verificar_pdf(
            caminho,
            usar_ocr=ocr,
            idioma=idioma,
            executavel_tesseract=executavel_tesseract,
            max_paginas_ocr=configuracao[0],
            tempo_total_ocr=configuracao[1],
            tempo_pagina_ocr=configuracao[2],
            dpi_ocr=configuracao[3],
            max_megapixels_ocr=configuracao[4],
            max_saida_ocr=configuracao[5],
        )
    if extensao == "docx":
        return _verificar_docx(caminho)
    return Verificacao(
        caminho, abriu=False, erro=f"extensão fora do escopo: .{extensao}"
    )


def sha256_arquivo(caminho: Path) -> str:
    digesto = hashlib.sha256()
    with caminho.open("rb") as arquivo:
        for bloco in iter(lambda: arquivo.read(1 << 20), b""):
            digesto.update(bloco)
    return digesto.hexdigest()
