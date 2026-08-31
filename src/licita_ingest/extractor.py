"""Extrator inicial ``documento → structured blocks`` da R3.

A implementação mantém duas propriedades importantes para as etapas seguintes:

* cada bloco carrega ``document_id`` e página (física no PDF, lógica no DOCX);
* o texto de parágrafos/células não é normalizado, resumido ou descartado.

PDFs sem camada textual usam o mesmo fallback OCR da coleta
(``licita_corpus.verify``). Sem texto utilizável, a extração falha de forma
explícita.
"""

from __future__ import annotations

import hashlib
from dataclasses import dataclass, replace
from pathlib import Path
from typing import Any, Iterable, Sequence

from .errors import OCRRequiredError, UnsupportedFormatError
from .models import (
    BBox,
    BlockType,
    DocumentFormat,
    PageKind,
    StructuredBlock,
    StructuredDocument,
    StructuredPage,
)


# ---------------------------------------------------------------------------
# General helpers


def sha256_file(path: str | Path, *, chunk_size: int = 1 << 20) -> str:
    """Calcula o SHA-256 do arquivo sem carregar tudo na memória."""

    digest = hashlib.sha256()
    with Path(path).open("rb") as file:
        for chunk in iter(lambda: file.read(chunk_size), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _document_id(path: Path, digest: str, requested: str | None) -> str:
    if requested is not None:
        if not requested.strip():
            raise ValueError("document_id não pode ser vazio")
        return requested
    # O hash evita que duas cópias de arquivos com o mesmo nome se confundam.
    # Um chamador que já tenha um id de processo deve fornecê-lo explicitamente.
    return f"sha256:{digest}"


def _path_and_identity(
    source: str | Path, document_id: str | None
) -> tuple[Path, str, str]:
    path = Path(source)
    if not path.exists():
        raise FileNotFoundError(path)
    if not path.is_file():
        raise IsADirectoryError(path)
    digest = sha256_file(path)
    return path, digest, _document_id(path, digest, document_id)


def _bbox(value: Any) -> BBox | None:
    if value is None:
        return None
    try:
        values = tuple(float(value[index]) for index in range(4))
    except (IndexError, KeyError, TypeError, ValueError):
        return None
    return values  # type: ignore[return-value]


def _bbox_area(box: BBox) -> float:
    return max(0.0, box[2] - box[0]) * max(0.0, box[3] - box[1])


def _overlap_ratio(inner: BBox, outer: BBox) -> float:
    """Retorna quanto da caixa ``inner`` está dentro de ``outer``."""

    left = max(inner[0], outer[0])
    top = max(inner[1], outer[1])
    right = min(inner[2], outer[2])
    bottom = min(inner[3], outer[3])
    intersection = max(0.0, right - left) * max(0.0, bottom - top)
    area = _bbox_area(inner)
    if area == 0:
        center_x = (inner[0] + inner[2]) / 2
        center_y = (inner[1] + inner[3]) / 2
        return float(
            outer[0] <= center_x <= outer[2]
            and outer[1] <= center_y <= outer[3]
        )
    return intersection / area


def _cluster(values: Iterable[float], tolerance: float = 1e-3) -> list[float]:
    """Agrupa coordenadas praticamente iguais, mantendo-as ordenadas."""

    clusters: list[float] = []
    for value in sorted(values):
        if not clusters or abs(value - clusters[-1]) > tolerance:
            clusters.append(value)
    return clusters


# ---------------------------------------------------------------------------
# PDF


@dataclass(frozen=True, slots=True)
class _PdfTextRecord:
    bbox: BBox
    text: str
    number: int
    block_type: int


def _pdf_text_records(page: Any) -> list[_PdfTextRecord]:
    records: list[_PdfTextRecord] = []
    for position, raw in enumerate(page.get_text("blocks", sort=False)):
        box = _bbox(raw)
        if box is None:
            continue
        if len(raw) <= 4:
            continue
        text = raw[4] if isinstance(raw[4], str) else str(raw[4])
        try:
            number = int(raw[5])
        except (IndexError, TypeError, ValueError):
            number = position
        try:
            block_type = int(raw[6])
        except (IndexError, TypeError, ValueError):
            block_type = 0
        records.append(_PdfTextRecord(box, text, number, block_type))
    return records


def _find_pdf_tables(page: Any) -> list[Any]:
    """Localiza tabelas sem tornar a leitura textual dependente delas.

    A detecção de tabela do PyMuPDF pode falhar em PDFs sem linhas ou variar
    entre versões. Nesse caso, os blocos textuais continuam sendo retornados;
    não há inferência silenciosa de uma tabela.
    """

    find_tables = getattr(page, "find_tables", None)
    if find_tables is None:
        return []
    try:
        finder = find_tables()
    except Exception:
        return []
    tables = getattr(finder, "tables", ())
    return list(tables or ())


def _table_matrix(table: Any) -> tuple[list[list[Any]], int, int]:
    try:
        extracted = table.extract()
    except Exception:
        extracted = []

    matrix: list[list[Any]] = []
    if isinstance(extracted, (list, tuple)):
        for row in extracted:
            if isinstance(row, (list, tuple)):
                matrix.append(list(row))

    try:
        row_count = int(table.row_count)
    except (AttributeError, TypeError, ValueError):
        row_count = len(matrix)
    try:
        column_count = int(table.col_count)
    except (AttributeError, TypeError, ValueError):
        column_count = max((len(row) for row in matrix), default=0)

    row_count = max(row_count, len(matrix))
    column_count = max(column_count, max((len(row) for row in matrix), default=0))
    return matrix, row_count, column_count


def _table_cell_boxes(table: Any, row_count: int, column_count: int) -> dict[tuple[int, int], BBox]:
    raw_boxes = [box for box in (_bbox(value) for value in (getattr(table, "cells", ()) or ())) if box]
    if not raw_boxes or row_count == 0 or column_count == 0:
        return {}

    row_starts = _cluster(box[1] for box in raw_boxes)
    column_starts = _cluster(box[0] for box in raw_boxes)
    if not row_starts or not column_starts:
        return {}

    result: dict[tuple[int, int], BBox] = {}
    for box in raw_boxes:
        row = min(range(row_count), key=lambda index: abs(row_starts[min(index, len(row_starts) - 1)] - box[1]))
        column = min(
            range(column_count),
            key=lambda index: abs(column_starts[min(index, len(column_starts) - 1)] - box[0]),
        )
        result.setdefault((row, column), box)
    return result


def _pdf_words_in_box(page: Any, box: BBox) -> str:
    """Fallback de célula quando a tabela não entrega uma matriz textual."""

    try:
        words = page.get_text("words", sort=False)
    except Exception:
        return ""
    selected: list[tuple[float, float, str]] = []
    for word in words:
        word_box = _bbox(word)
        if word_box is None or _overlap_ratio(word_box, box) < 0.5:
            continue
        try:
            text = str(word[4])
        except (IndexError, TypeError):
            continue
        selected.append((word_box[1], word_box[0], text))
    selected.sort(key=lambda item: (item[0], item[1]))
    return " ".join(item[2] for item in selected)


def _pdf_cell_text(
    page: Any,
    records: Sequence[_PdfTextRecord],
    matrix: Sequence[Sequence[Any]],
    row: int,
    column: int,
    box: BBox | None,
) -> str:
    # Quando um bloco de texto cabe integralmente na célula, mantê-lo sem
    # aparar é a forma mais fiel de conservar as quebras retornadas pelo PDF.
    if box is not None:
        contained = [
            record
            for record in records
            if record.block_type == 0 and _overlap_ratio(record.bbox, box) >= 0.999
        ]
        if contained:
            contained.sort(key=lambda item: (item.bbox[1], item.bbox[0], item.number))
            return "".join(record.text for record in contained)

    if row < len(matrix) and column < len(matrix[row]):
        value = matrix[row][column]
        if value is not None and str(value) != "":
            return str(value)

    if box is not None:
        return _pdf_words_in_box(page, box)
    return ""


def _pdf_table_block(
    page: Any,
    document_id: str,
    page_number: int,
    table_number: int,
    table: Any,
    records: Sequence[_PdfTextRecord],
) -> StructuredBlock:
    table_id = f"{document_id}:p-{page_number:04d}:t-{table_number:04d}"
    table_box = _bbox(getattr(table, "bbox", None))
    matrix, row_count, column_count = _table_matrix(table)
    boxes = _table_cell_boxes(table, row_count, column_count)

    cells: list[StructuredBlock] = []
    rows_as_text: list[str] = []
    for row in range(row_count):
        row_cells: list[str] = []
        for column in range(column_count):
            cell_box = boxes.get((row, column))
            cell_id = f"{table_id}:r-{row + 1:04d}:c-{column + 1:04d}"
            cell_text = _pdf_cell_text(
                page, records, matrix, row, column, cell_box
            )
            row_cells.append(cell_text)
            cells.append(
                StructuredBlock(
                    document_id=document_id,
                    id=cell_id,
                    type=BlockType.TABLE_CELL,
                    text=cell_text,
                    page=page_number,
                    page_kind=PageKind.PHYSICAL,
                    index=row * column_count + column,
                    table_index=table_number,
                    row_index=row,
                    column_index=column,
                    parent_id=table_id,
                    bbox=cell_box,
                    metadata={"source": "pymupdf.table_cell"},
                )
            )
        rows_as_text.append("\t".join(row_cells))

    return StructuredBlock(
        document_id=document_id,
        id=table_id,
        type=BlockType.TABLE,
        text="\n".join(rows_as_text),
        page=page_number,
        page_kind=PageKind.PHYSICAL,
        index=table_number,
        table_index=table_number,
        bbox=table_box,
        children=tuple(cells),
        metadata={
            "source": "pymupdf.find_tables",
            "row_count": row_count,
            "column_count": column_count,
            "table_text_is_cell_aggregation": True,
        },
    )


def _pdf_page_blocks(
    page: Any, document_id: str, page_number: int
) -> tuple[StructuredBlock, ...]:
    records = _pdf_text_records(page)
    tables = _find_pdf_tables(page)
    table_blocks = [
        _pdf_table_block(page, document_id, page_number, number, table, records)
        for number, table in enumerate(tables, start=1)
    ]
    table_boxes = [block.bbox for block in table_blocks if block.bbox is not None]

    candidates: list[tuple[float, float, int, StructuredBlock]] = []
    sequence = 0
    for record in records:
        # A table gets one structured block and cells. Do not expose the same
        # table row a second time as an ordinary paragraph block.
        if record.block_type == 0 and any(
            _overlap_ratio(record.bbox, table_box) >= 0.5 for table_box in table_boxes
        ):
            continue
        block_type = BlockType.PARAGRAPH if record.block_type == 0 else BlockType.IMAGE
        block_id = f"{document_id}:p-{page_number:04d}:b-{record.number + 1:04d}"
        block = StructuredBlock(
            document_id=document_id,
            id=block_id,
            type=block_type,
            text=record.text if block_type is BlockType.PARAGRAPH else "",
            page=page_number,
            page_kind=PageKind.PHYSICAL,
            index=record.number,
            bbox=record.bbox,
            metadata={
                "source": "pymupdf.text_block"
                if block_type is BlockType.PARAGRAPH
                else "pymupdf.image_block",
                "pdf_block_type": record.block_type,
            },
        )
        candidates.append((record.bbox[1], record.bbox[0], sequence, block))
        sequence += 1

    for table_block in table_blocks:
        if table_block.bbox is None:
            position = (float("inf"), float("inf"))
        else:
            position = (table_block.bbox[1], table_block.bbox[0])
        candidates.append((position[0], position[1], sequence, table_block))
        sequence += 1

    candidates.sort(key=lambda item: (item[0], item[1], item[2]))
    return tuple(replace(block, index=index) for index, (_, _, _, block) in enumerate(candidates))


def _pdf_has_unreadable_image(page: Any, page_text: str) -> bool:
    if page_text.strip():
        return False
    try:
        return bool(page.get_images(full=True))
    except Exception:
        return False


def _aplicar_ocr_pdf(
    path: Path,
    document_id: str,
    pages: Sequence[StructuredPage],
    ocr_pages: Sequence[int],
) -> tuple[tuple[StructuredPage, ...], tuple[int, ...]]:
    """Reusa o motor de OCR da coleta; não inventa texto se a qualidade falhar."""
    from licita_corpus.verify import verificar

    resultado = verificar(path, ocr=True)
    textos = {
        pagina.pagina: pagina.texto
        for pagina in resultado.paginas_avaliadas
        if pagina.texto.strip()
    }
    atualizadas: list[StructuredPage] = []
    restantes: list[int] = []
    alvo = set(ocr_pages)
    for page in pages:
        if page.number not in alvo:
            atualizadas.append(page)
            continue
        texto = textos.get(page.number, "").strip()
        if not texto:
            restantes.append(page.number)
            atualizadas.append(page)
            continue
        bloco = StructuredBlock(
            document_id=document_id,
            id=f"{document_id}:p-{page.number:04d}:ocr-0001",
            type=BlockType.PARAGRAPH,
            text=texto,
            page=page.number,
            page_kind=PageKind.PHYSICAL,
            index=len(page.blocks),
            metadata={"source": "ocr"},
        )
        atualizadas.append(replace(page, text=texto, blocks=page.blocks + (bloco,)))
    return tuple(atualizadas), tuple(restantes)


def _extract_pdf(path: Path, document_id: str, digest: str) -> StructuredDocument:
    import pymupdf

    pages: list[StructuredPage] = []
    ocr_pages: list[int] = []
    with pymupdf.open(str(path)) as pdf:
        for number, page in enumerate(pdf, start=1):
            page_text = page.get_text()
            blocks = _pdf_page_blocks(page, document_id, number)
            if _pdf_has_unreadable_image(page, page_text):
                ocr_pages.append(number)
            pages.append(
                StructuredPage(
                    document_id=document_id,
                    number=number,
                    text=page_text,
                    blocks=blocks,
                )
            )
        page_count = pdf.page_count

    if ocr_pages:
        pages, ocr_pages = _aplicar_ocr_pdf(path, document_id, pages, ocr_pages)
        if ocr_pages:
            raise OCRRequiredError(path, ocr_pages)

    return StructuredDocument(
        document_id=document_id,
        format=DocumentFormat.PDF,
        source=str(path),
        sha256=digest,
        blocks=tuple(block for page in pages for block in page.blocks),
        pages=tuple(pages),
        page_count=page_count,
        page_kind=PageKind.PHYSICAL,
        ocr_required_pages=(),
    )


# ---------------------------------------------------------------------------
# DOCX


def _docx_style_name(paragraph: Any) -> str | None:
    try:
        return paragraph.style.name
    except (AttributeError, KeyError, ValueError):
        return None


def _has_explicit_page_break(paragraph: Any) -> bool:
    """Detecta quebras explícitas, sem fingir conhecer o layout do Word."""

    try:
        from docx.oxml.ns import qn

        for element in paragraph._p.iter():
            if element.tag == qn("w:lastRenderedPageBreak"):
                return True
            if element.tag == qn("w:br") and element.get(qn("w:type")) == "page":
                return True
    except (AttributeError, KeyError, ValueError):
        return False
    return False


@dataclass(slots=True)
class _DocxBuilder:
    document_id: str
    logical_page: int = 1
    paragraph_index: int = 0
    table_index: int = 0

    def paragraph_block(
        self,
        paragraph: Any,
        block_id: str,
        *,
        index: int,
        parent_id: str | None = None,
    ) -> StructuredBlock:
        paragraph_number = self.paragraph_index
        self.paragraph_index += 1
        return StructuredBlock(
            document_id=self.document_id,
            id=block_id,
            type=BlockType.PARAGRAPH,
            text=paragraph.text,
            page=self.logical_page,
            page_kind=PageKind.LOGICAL,
            index=index,
            paragraph_index=paragraph_number,
            parent_id=parent_id,
            style_name=_docx_style_name(paragraph),
            metadata={"source": "python-docx.paragraph"},
        )

    def table_block(self, table: Any, block_id: str, *, index: int) -> StructuredBlock:
        table_number = self.table_index
        self.table_index += 1
        cells: list[StructuredBlock] = []
        rows_as_text: list[str] = []

        for row_number, row in enumerate(table.rows):
            row_text: list[str] = []
            for column_number, cell in enumerate(row.cells):
                cell_id = f"{block_id}:r-{row_number + 1:04d}:c-{column_number + 1:04d}"
                paragraphs: list[StructuredBlock] = []
                for paragraph_number, paragraph in enumerate(cell.paragraphs):
                    paragraph_id = (
                        f"{cell_id}:p-{paragraph_number + 1:04d}"
                    )
                    paragraphs.append(
                        self.paragraph_block(
                            paragraph,
                            paragraph_id,
                            index=paragraph_number,
                            parent_id=cell_id,
                        )
                    )
                cell_text = cell.text
                row_text.append(cell_text)
                cells.append(
                    StructuredBlock(
                        document_id=self.document_id,
                        id=cell_id,
                        type=BlockType.TABLE_CELL,
                        text=cell_text,
                        page=self.logical_page,
                        page_kind=PageKind.LOGICAL,
                        index=row_number * len(row.cells) + column_number,
                        table_index=table_number,
                        row_index=row_number,
                        column_index=column_number,
                        parent_id=block_id,
                        children=tuple(paragraphs),
                        metadata={"source": "python-docx.table_cell"},
                    )
                )
            rows_as_text.append("\t".join(row_text))

        return StructuredBlock(
            document_id=self.document_id,
            id=block_id,
            type=BlockType.TABLE,
            text="\n".join(rows_as_text),
            page=self.logical_page,
            page_kind=PageKind.LOGICAL,
            index=index,
            table_index=table_number,
            children=tuple(cells),
            metadata={
                "source": "python-docx.table",
                "row_count": len(table.rows),
                "column_count": max((len(row.cells) for row in table.rows), default=0),
                "table_text_is_cell_aggregation": True,
            },
        )


def _docx_body_children(document: Any) -> Iterable[tuple[str, Any]]:
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield "table", Table(child, document)


def _extract_docx(path: Path, document_id: str, digest: str) -> StructuredDocument:
    import docx

    document = docx.Document(str(path))
    builder = _DocxBuilder(document_id=document_id)
    blocks: list[StructuredBlock] = []

    for body_index, (kind, value) in enumerate(_docx_body_children(document)):
        if kind == "paragraph":
            block = builder.paragraph_block(
                value,
                f"{document_id}:p-{builder.paragraph_index + 1:04d}",
                index=body_index,
            )
            blocks.append(block)
            if _has_explicit_page_break(value):
                builder.logical_page += 1
        else:
            blocks.append(
                builder.table_block(
                    value,
                    f"{document_id}:t-{builder.table_index + 1:04d}",
                    index=body_index,
                )
            )

    logical_page_count = max(
        (block.page or 0 for block in blocks),
        default=0,
    )
    return StructuredDocument(
        document_id=document_id,
        format=DocumentFormat.DOCX,
        source=str(path),
        sha256=digest,
        blocks=tuple(blocks),
        pages=(),
        page_count=None,
        page_kind=PageKind.LOGICAL,
        logical_page_count=logical_page_count,
    )


# ---------------------------------------------------------------------------
# Public API


def extract_pdf(
    source: str | Path, *, document_id: str | None = None
) -> StructuredDocument:
    """Extrai texto, blocos e tabelas de um PDF usando PyMuPDF."""

    path, digest, resolved_id = _path_and_identity(source, document_id)
    if path.suffix.lower() != ".pdf":
        raise UnsupportedFormatError(f"arquivo não é PDF: {path}")
    return _extract_pdf(path, resolved_id, digest)


def extract_docx(
    source: str | Path, *, document_id: str | None = None
) -> StructuredDocument:
    """Extrai parágrafos e tabelas de um DOCX usando python-docx."""

    path, digest, resolved_id = _path_and_identity(source, document_id)
    if path.suffix.lower() != ".docx":
        raise UnsupportedFormatError(f"arquivo não é DOCX: {path}")
    return _extract_docx(path, resolved_id, digest)


def extract_document(
    source: str | Path, *, document_id: str | None = None
) -> StructuredDocument:
    """Seleciona o leitor pela extensão e retorna blocos estruturados."""

    path, digest, resolved_id = _path_and_identity(source, document_id)
    extension = path.suffix.lower()
    if extension == ".pdf":
        return _extract_pdf(path, resolved_id, digest)
    if extension == ".docx":
        return _extract_docx(path, resolved_id, digest)
    raise UnsupportedFormatError(
        f"extensão fora do escopo da R3: {path.suffix or '<sem extensão>'}; "
        "formatos aceitos: .pdf e .docx"
    )


# Alias curto para integrações que tratam o extrator como uma função genérica.
extract = extract_document
extrair_documento = extract_document


@dataclass(frozen=True, slots=True)
class DocumentExtractor:
    """Pequena fachada orientada a objeto para o mesmo extrator determinístico."""

    def extract(
        self, source: str | Path, *, document_id: str | None = None
    ) -> StructuredDocument:
        return extract_document(source, document_id=document_id)


__all__ = [
    "DocumentExtractor",
    "extract",
    "extract_docx",
    "extract_document",
    "extract_pdf",
    "extrair_documento",
    "sha256_file",
]
