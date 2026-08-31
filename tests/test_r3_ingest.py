"""Testes do primeiro extrator R3 com documentos gerados em tempo de teste."""

from __future__ import annotations

import hashlib
from pathlib import Path

import pytest
import pymupdf
from docx import Document as DocxDocument

from licita_ingest import (
    BlockType,
    DocumentFormat,
    OCRRequiredError,
    OCR_FALLBACK_IMPLEMENTED,
    PageKind,
    UnsupportedFormatError,
    extract_document,
    sha256_file,
)


def _make_synthetic_pdf(path: Path) -> str:
    document = pymupdf.open()
    page = document.new_page(width=480, height=300)
    page.insert_text(
        (40, 42),
        "Trecho PDF original — não normalizar.",
        fontsize=11,
    )

    x_coordinates = (40, 180, 320, 440)
    y_coordinates = (75, 115, 155)
    for x in x_coordinates:
        page.draw_line((x, y_coordinates[0]), (x, y_coordinates[-1]))
    for y in y_coordinates:
        page.draw_line((x_coordinates[0], y), (x_coordinates[-1], y))

    values = (
        ((50, 103), "Item"),
        ((190, 103), "Quantidade"),
        ((330, 103), "Unidade"),
        ((50, 143), "Caneta azul"),
        ((190, 143), "12"),
        ((330, 143), "un"),
    )
    for position, text in values:
        page.insert_text(position, text, fontsize=10)

    document.save(path)
    document.close()

    with pymupdf.open(path) as reopened:
        return reopened[0].get_text()


def _make_synthetic_docx(path: Path) -> None:
    document = DocxDocument()
    document.add_paragraph("  Texto DOCX original — com acentuação.  ")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Quantidade"
    table.cell(1, 0).text = "Caderno"
    table.cell(1, 1).text = "7"
    document.add_page_break()
    document.add_paragraph("Parágrafo da segunda página lógica")
    document.save(path)


def test_pdf_preserves_document_page_blocks_table_cells_and_original_text(
    tmp_path: Path,
) -> None:
    path = tmp_path / "sintetico.pdf"
    expected_page_text = _make_synthetic_pdf(path)

    extracted = extract_document(path, document_id="processo-1#tr")

    assert extracted.document_id == "processo-1#tr"
    assert extracted.format is DocumentFormat.PDF
    assert extracted.sha256 == hashlib.sha256(path.read_bytes()).hexdigest()
    assert sha256_file(path) == extracted.sha256
    assert extracted.page_count == 1
    assert extracted.page_kind is PageKind.PHYSICAL
    assert extracted.pages[0].number == 1
    assert extracted.pages[0].text == expected_page_text

    paragraphs = [block for block in extracted.blocks if block.type is BlockType.PARAGRAPH]
    assert len(paragraphs) == 1
    assert paragraphs[0].document_id == extracted.document_id
    assert paragraphs[0].page == 1
    # A codificação literal do PDF pode mapear um glifo para outro caractere
    # no texto extraído; o contrato da R3 é preservar exatamente o retorno do
    # leitor, sem normalizar depois dele.
    assert paragraphs[0].text == expected_page_text.splitlines(keepends=True)[0]

    tables = extracted.tables
    assert len(tables) == 1
    table = tables[0]
    assert table.page == 1
    assert table.table_index == 1
    cells = {(cell.row_index, cell.column_index): cell for cell in table.cells}
    assert cells[(0, 0)].text == "Item"
    assert cells[(0, 1)].text == "Quantidade"
    assert cells[(1, 0)].text == "Caneta azul"
    assert cells[(1, 1)].text == "12"
    assert cells[(1, 2)].text == "un"
    assert table.text.splitlines()[1].startswith("Caneta azul")

    all_ids = [block.id for block in extracted.iter_blocks()]
    assert len(all_ids) == len(set(all_ids))
    assert extracted.get_block(cells[(1, 1)].id).text == "12"
    assert all(block.document_id == extracted.document_id for block in extracted.iter_blocks())


def test_docx_preserves_paragraphs_tables_and_explicit_logical_page_break(
    tmp_path: Path,
) -> None:
    path = tmp_path / "sintetico.docx"
    _make_synthetic_docx(path)

    extracted = extract_document(path, document_id="docx-sintetico")

    assert extracted.document_id == "docx-sintetico"
    assert extracted.format is DocumentFormat.DOCX
    assert extracted.sha256 == hashlib.sha256(path.read_bytes()).hexdigest()
    assert extracted.pages == ()
    assert extracted.page_count is None
    assert extracted.page_kind is PageKind.LOGICAL
    assert extracted.logical_page_count == 2

    top_level = list(extracted.iter_blocks(include_children=False))
    assert [block.type for block in top_level] == [
        BlockType.PARAGRAPH,
        BlockType.TABLE,
        BlockType.PARAGRAPH,
        BlockType.PARAGRAPH,
    ]
    assert top_level[0].text == "  Texto DOCX original — com acentuação.  "
    assert top_level[0].page == 1
    assert top_level[-1].text == "Parágrafo da segunda página lógica"
    assert top_level[-1].page == 2

    table = extracted.tables[0]
    assert table.page == 1
    cells = {(cell.row_index, cell.column_index): cell for cell in table.cells}
    assert cells[(0, 0)].text == "Item"
    assert cells[(0, 1)].text == "Quantidade"
    assert cells[(1, 0)].text == "Caderno"
    assert cells[(1, 1)].text == "7"
    assert cells[(0, 0)].children[0].type is BlockType.PARAGRAPH
    assert cells[(0, 0)].children[0].text == "Item"
    assert all(block.document_id == extracted.document_id for block in extracted.iter_blocks())


def test_scanned_pdf_reports_failed_ocr_instead_of_silently_returning_empty_text(
    tmp_path: Path,
) -> None:
    path = tmp_path / "escaneado.pdf"
    document = pymupdf.open()
    page = document.new_page(width=100, height=100)
    pixmap = pymupdf.Pixmap(pymupdf.csRGB, (0, 0, 40, 40), False)
    pixmap.clear_with(255)
    page.insert_image((20, 20, 60, 60), stream=pixmap.tobytes("png"))
    document.save(path)
    document.close()

    assert OCR_FALLBACK_IMPLEMENTED is True
    with pytest.raises(OCRRequiredError, match="OCR não produziu texto utilizável") as error:
        extract_document(path, document_id="scan-1")
    assert error.value.pages == (1,)


def test_unsupported_extension_is_explicit(tmp_path: Path) -> None:
    path = tmp_path / "entrada.txt"
    path.write_text("não é um documento suportado", encoding="utf-8")

    with pytest.raises(UnsupportedFormatError, match="formatos aceitos"):
        extract_document(path)
